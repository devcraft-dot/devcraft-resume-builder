"""Build .docx files in memory (BytesIO) for direct Google Drive upload."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

from app.services.resume_docx_formatter import build_formatted_resume_docx

_BODY_FONT = "Calibri"


@dataclass
class JobData:
    """Lightweight stand-in for what the docx formatters expect via getattr."""
    title: str = ""
    company_name: str = ""
    description_text: str = ""
    url: str = ""
    fields: list | None = None


def _safe_filename(value: str) -> str:
    return re.sub(r"[^\w\s-]", "", str(value or "")).strip().replace(" ", "_")


def _doc_to_buffer(doc: Document) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _apply_standard_page_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0


def _strip_default_paragraph(doc: Document) -> None:
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)


# ---------------------------------------------------------------------------
# Answers Q&A parsing
# ---------------------------------------------------------------------------

def _looks_like_question_line(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if re.match(r"^#{1,3}\s", s):
        return True
    if re.match(r"^\*\*(.+?)\*\*\s*$", s):
        return True
    if re.match(r"^(?:Q|Question)\s*[:.]?\s*\S", s, re.IGNORECASE):
        return True
    if re.match(r"^\d+\.\s+\S", s):
        return True
    if re.match(r"^[-*]\s+\S", s):
        return True
    return False


def _parse_answers_to_qa_pairs(text: str) -> list[tuple[str, str]]:
    text = str(text or "").strip()
    if not text:
        return []
    text = re.sub(
        r"^#{1,3}\s*Application\s+Questions\s*\n*",
        "", text, count=1, flags=re.IGNORECASE | re.MULTILINE,
    ).strip()

    pairs: list[tuple[str, str]] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        m = re.match(r"^\*\*(.+?)\*\*\s*$", line)
        if m:
            q = m.group(1).strip()
            ans_lines: list[str] = []
            i += 1
            while i < len(lines):
                L = lines[i].strip()
                if not L:
                    i += 1
                    if ans_lines and i < len(lines) and _looks_like_question_line(lines[i]):
                        break
                    continue
                if _looks_like_question_line(lines[i]):
                    break
                ans_lines.append(L)
                i += 1
            pairs.append((q, "\n".join(ans_lines).strip()))
            continue

        m = re.match(r"^(?:Q|Question)\s*[:.]?\s*(.+)$", line, re.IGNORECASE)
        if m:
            q = m.group(1).strip()
            ans_lines = []
            i += 1
            while i < len(lines):
                L = lines[i].strip()
                if not L:
                    i += 1
                    if ans_lines and i < len(lines) and _looks_like_question_line(lines[i]):
                        break
                    continue
                if _looks_like_question_line(lines[i]):
                    break
                if re.match(r"^(?:A|Answer)\s*[:.]?\s*", L, re.IGNORECASE):
                    L = re.sub(r"^(?:A|Answer)\s*[:.]?\s*", "", L, flags=re.IGNORECASE)
                ans_lines.append(L)
                i += 1
            pairs.append((q, "\n".join(ans_lines).strip()))
            continue

        m = re.match(r"^\d+\.\s+(.+)$", line)
        if m:
            q = m.group(1).strip()
            ans_lines = []
            i += 1
            while i < len(lines):
                L = lines[i].strip()
                if not L:
                    i += 1
                    continue
                if re.match(r"^\d+\.\s+", L) or _looks_like_question_line(lines[i]):
                    break
                if re.match(r"^(?:A|Answer)\s*[:.]?\s*", L, re.IGNORECASE):
                    L = re.sub(r"^(?:A|Answer)\s*[:.]?\s*", "", L, flags=re.IGNORECASE)
                ans_lines.append(L)
                i += 1
            pairs.append((q, "\n".join(ans_lines).strip()))
            continue

        i += 1

    if pairs:
        return pairs

    blocks = [b.strip() for b in re.split(r"\n\s*\n+", text) if b.strip()]
    for b in blocks:
        ls = [x.strip() for x in b.splitlines() if x.strip()]
        if not ls:
            continue
        if len(ls) == 1:
            pairs.append(("Application question", ls[0]))
        else:
            pairs.append((ls[0], "\n".join(ls[1:]).strip()))

    if not pairs and text:
        pairs.append(("Application responses", text))

    return pairs


def _add_qa_pair_to_doc(doc: Document, question: str, answer: str) -> None:
    p_q = doc.add_paragraph()
    r_label = p_q.add_run("Question: ")
    r_label.bold = True
    r_label.font.name = _BODY_FONT
    r_label.font.size = Pt(10.5)
    r_q = p_q.add_run(question)
    r_q.font.name = _BODY_FONT
    r_q.font.size = Pt(10.5)
    p_q.paragraph_format.space_after = Pt(2)

    ans = (answer or "").strip()
    if not ans:
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.paragraph_format.space_after = Pt(10)
        ra = p_a.add_run("Answer: ")
        ra.bold = True
        ra.font.name = _BODY_FONT
        ra.font.size = Pt(10.5)
        ra2 = p_a.add_run("(No answer provided.)")
        ra2.italic = True
        ra2.font.name = _BODY_FONT
        ra2.font.size = Pt(10.5)
        return

    ans_lines = [x.strip() for x in ans.splitlines() if x.strip()] or [ans]
    first = True
    for segment in ans_lines:
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.paragraph_format.space_after = Pt(4 if first else 2)
        if first:
            ra = p_a.add_run("Answer: ")
            ra.bold = True
            ra.font.name = _BODY_FONT
            ra.font.size = Pt(10.5)
            rb = p_a.add_run(segment)
            rb.font.name = _BODY_FONT
            rb.font.size = Pt(10.5)
            first = False
        else:
            rb = p_a.add_run(segment)
            rb.font.name = _BODY_FONT
            rb.font.size = Pt(10.5)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# Public builders — all return (BytesIO, filename)
# ---------------------------------------------------------------------------

def build_resume_docx(resume_text: str, title: str, company_name: str,
                      description_text: str, profile_name: str) -> tuple[BytesIO, str]:
    job = JobData(title=title, company_name=company_name, description_text=description_text)
    return build_formatted_resume_docx(resume_text, job, profile_name)


def build_jd_docx(title: str, company_name: str, description_text: str) -> tuple[BytesIO, str]:
    doc = Document()
    _strip_default_paragraph(doc)
    _apply_standard_page_style(doc)

    title_para = doc.add_paragraph()
    run = title_para.add_run(f"Job Description: {title} @ {company_name}".strip())
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(14)

    for line in str(description_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("##") or stripped.startswith("==") or stripped.isupper():
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip("#= ").strip())
            r.bold = True
            r.font.name = _BODY_FONT
            r.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph(stripped)
            for r in p.runs:
                r.font.name = _BODY_FONT
                r.font.size = Pt(10.5)

    filename = f"JD_{_safe_filename(title)}_{date.today().isoformat()}.docx"
    return _doc_to_buffer(doc), filename


def _split_cover_letter_prefix(body: str) -> tuple[str, str]:
    """If body starts with ## Cover Letter, return (cover_plain, rest for Q&A parser)."""
    b = str(body or "").strip()
    if not re.match(r"^##\s*Cover\s+Letter\s*$", b.split("\n", 1)[0].strip(), re.I):
        return "", b
    lines = b.split("\n", 1)
    if len(lines) < 2:
        return "", b
    rest = lines[1].lstrip("\n")
    m = re.search(r"^\s*##\s*Application\s+Questions\s*$", rest, re.I | re.M)
    if m:
        cover = rest[: m.start()].strip()
        qa = rest[m.end() :].lstrip("\n")
        return cover, qa
    return rest.strip(), ""


def build_answers_docx(title: str, company_name: str, answers_text: str) -> tuple[BytesIO, str]:
    doc = Document()
    _strip_default_paragraph(doc)
    _apply_standard_page_style(doc)

    title_para = doc.add_paragraph()
    run = title_para.add_run(f"Application packet: {title} @ {company_name}")
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(14)
    title_para.paragraph_format.space_after = Pt(8)

    body = str(answers_text or "").strip()
    if body:
        cover, qa_source = _split_cover_letter_prefix(body)
        if cover:
            h = doc.add_paragraph()
            hr = h.add_run("Cover Letter")
            hr.bold = True
            hr.font.name = _BODY_FONT
            hr.font.size = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            for block in re.split(r"\n\s*\n+", cover):
                para = block.strip()
                if not para:
                    continue
                p = doc.add_paragraph()
                r = p.add_run(para)
                r.font.name = _BODY_FONT
                r.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(6)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(10)

        pairs = _parse_answers_to_qa_pairs(qa_source)
        for q, a in pairs:
            _add_qa_pair_to_doc(doc, q, a)
        if not cover and not pairs:
            p = doc.add_paragraph()
            r = p.add_run("No application Q&A block was parsed (optional for this job).")
            r.font.name = _BODY_FONT
            r.font.size = Pt(10.5)
    else:
        p = doc.add_paragraph()
        r = p.add_run("No application answers were generated for this job.")
        r.font.name = _BODY_FONT
        r.font.size = Pt(10.5)

    filename = f"Answers_{_safe_filename(title)}_{date.today().isoformat()}.docx"
    return _doc_to_buffer(doc), filename
