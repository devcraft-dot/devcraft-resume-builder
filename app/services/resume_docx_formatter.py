"""
ATS-oriented resume .docx generation.

Renders model Markdown into .docx: layout, sections, bullets, and ``**emphasis**`` runs
exactly as written. Structural styling only (name, section titles, experience header rows,
Skills category labels). Skill list values are forced plain so only the category label is bold.

Layout follows ATS + recruiter guidance (e.g. [HireKit resume formatting](https://hirekit.co/guides/resume-formatting-best-practices), [ATS resume fonts & margins](https://atsresumeai.com/blog/ats-resume-formatting-guide/)): **single column**, standard **Calibri** body **~11 pt**, margins **0.9 in** all sides (within the usual **0.75–1 in** band; avoid going under about **0.5 in**), **~1.15** line spacing for scan speed, **bold section titles**. **Name + contact (optional title) centered**; everything else **left-aligned**. No critical text in Word headers/footers.

For quick scans, bold **job title and company** on one line creates clear anchors ([Resumly section scanning](https://www.resumly.ai/blog/optimizing-resume-sections-for-quick-scanning-by-recruiters)).

**Bullets (Experience, etc.):** **0.25 in** symmetric hanging indent (Word-style default depth). **Space after** bullets stays in a **~6–9 pt** skim band for short-to-medium bullets ([HireKit spacing notes](https://hirekit.co/guides/resume-formatting-best-practices)). Leader is **•** plus a normal space (plain text, ATS-safe).

**Skills (Category: values):** each line uses a **left tab stop** after the bold label so wrapped
value text aligns under the value column—the usual Word pattern for label/value lines and
recommended for consistent alignment in structured sections ([tab stops / ruler alignment](https://www.therecruitmentinsights.com/blog/resume-writing-tipstricks/005-resume-formatting-in-ms-wordstep-by-step)).

Light touches on text: curly quotes to ASCII, collapsed whitespace, and date tokens in
lines parsed through ``_clean_line`` (role/education rows, body lines) for consistent
``Month YYYY`` spelling and **en-dash** (U+2013) between range endpoints — no verb substitution,
JD keyword injection, or spelling rewrites.

Known sections are reordered for output to: Summary → Skills → Experience (+ projects) →
Education → other sections (e.g. certifications, application Q&A), regardless of model order.
"""

from __future__ import annotations

import re
from datetime import date
from io import BytesIO

# ── Section detection (same semantics as automation) ──────────────────────────
_SECTION_KEYWORDS = {
    "SUMMARY",
    "MARKET TITLE",
    "PROFESSIONAL SUMMARY",
    "OBJECTIVE",
    "SKILLS",
    "TECHNICAL SKILLS",
    "CORE COMPETENCIES",
    "WORK EXPERIENCE",
    "PROFESSIONAL EXPERIENCE",
    "EXPERIENCE",
    "PROJECTS",
    "NOTABLE PROJECTS",
    "KEY PROJECTS",
    "EDUCATION",
    "CERTIFICATIONS",
    "CERTIFICATIONS OR ACHIEVEMENTS",
    "ACHIEVEMENTS",
    "AWARDS",
    "ADDITIONAL INFORMATION",
    "APPLICATION QUESTIONS",
}

_EXPERIENCE_SECTIONS = {
    "WORK EXPERIENCE",
    "PROFESSIONAL EXPERIENCE",
    "EXPERIENCE",
    "EDUCATION",
}

# Role + bullet blocks (exclude EDUCATION so education bullets stay a bit tighter).
_EXPERIENCE_ROLE_SECTIONS = frozenset(
    {
        "WORK EXPERIENCE",
        "PROFESSIONAL EXPERIENCE",
        "EXPERIENCE",
        "PROJECTS",
        "NOTABLE PROJECTS",
        "KEY PROJECTS",
    }
)

# System fonts ATS parsers handle well; 11 pt body is widely recommended.
_BODY_FONT = "Calibri"
_BODY_PT = 11.0
_NAME_PT = 18.0  # guides often cite ~14–18 pt for name (bold)
_SECTION_TITLE_PT = 13.0  # section titles commonly 12–14 pt bold
_LINE_SPACING = 1.15  # single-ish with slight air; common “readable default” band
_MARGINS_IN = 0.9  # 0.75–1 in is the usual safe band; 0.9 balances text width + whitespace
# Hanging bullet: match Word’s common ~0.25 in list hang ([Microsoft hanging indent](https://support.microsoft.com/en-us/office/indent-the-second-line-in-word-9d1b9955-d08a-4773-a900-d0a9e641279c)).
_BULLET_LEFT_INDENT_IN = 0.25
_BULLET_FIRST_LINE_INDENT_IN = -0.25
_BULLET_LEADER = "\u2022 "
# Skills: label + tab + values; tab position balances short vs long category labels.
_SKILL_VALUE_TAB_IN = 1.78


def _safe_filename(value: str) -> str:
    return re.sub(r"[^\w\s-]", "", str(value or "")).strip().replace(" ", "_")


def _strip_md_bold(text: str) -> str:
    s = text.strip()
    if s.startswith("**") and s.endswith("**") and len(s) > 4:
        inner = s[2:-2].strip()
        if "**" not in inner:
            return inner
    return s


def _strip_md_spans_for_skill_values(text: str) -> str:
    """Skills: strip ``**...**`` so only the formatter's category label stays bold."""
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", str(text or ""))


def _strip_md_heading(text: str) -> str:
    return re.sub(r"^#{1,6}\s*", "", text.strip())


def _sanitize_text(text: str) -> str:
    """Normalize encoding-safe whitespace and quotes; does not rewrite vocabulary."""
    s = text
    s = s.replace("\u2018", "'").replace("\u2019", "'")
    s = s.replace("\u201c", '"').replace("\u201d", '"')
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


_MONTH_ABBR_TO_FULL = {
    "jan": "January",
    "feb": "February",
    "mar": "March",
    "apr": "April",
    "may": "May",
    "jun": "June",
    "jul": "July",
    "aug": "August",
    "sep": "September",
    "sept": "September",
    "oct": "October",
    "nov": "November",
    "dec": "December",
}
_FULL_MONTH_RE = re.compile(
    r"\b("
    + "|".join(
        [
            "january",
            "february",
            "march",
            "april",
            "may",
            "june",
            "july",
            "august",
            "september",
            "october",
            "november",
            "december",
        ]
    )
    + r")\s+(\d{4})\b",
    re.IGNORECASE,
)
_ABBR_MONTH_RE = re.compile(
    r"\b(" + "|".join(_MONTH_ABBR_TO_FULL) + r")\.?\s+(\d{4})\b",
    re.IGNORECASE,
)


def _normalize_dates(text: str) -> str:
    """Normalize dates to full month names and a simple hyphen separator."""
    def _full(m: re.Match) -> str:
        return f"{m.group(1).capitalize()} {m.group(2)}"

    def _expand(m: re.Match) -> str:
        return f"{_MONTH_ABBR_TO_FULL[m.group(1).lower().rstrip('.')]} {m.group(2)}"

    s = _FULL_MONTH_RE.sub(_full, text)
    s = _ABBR_MONTH_RE.sub(_expand, s)
    # Standardize range punctuation to en dash (May 2025 – March 2026).
    en = "\u2013"

    def _range_dash(m: re.Match) -> str:
        return f"{m.group(1)} {en} {m.group(2)}"

    s = re.sub(
        r"([A-Za-z]+\s+\d{4})\s*[–—-]\s*([A-Za-z]+\s+\d{4}|Present)",
        _range_dash,
        s,
    )
    return s


def _clean_line(text: str) -> str:
    return _normalize_dates(_sanitize_text(_strip_md_bold(_strip_md_heading(text))))


def _is_section_header(line: str) -> bool:
    raw = line.strip()
    stripped = _clean_line(raw)
    head = _strip_md_heading(raw)
    if head.upper() in _SECTION_KEYWORDS or stripped.upper() in _SECTION_KEYWORDS:
        return True
    check = head.strip()
    return (
        check.isupper()
        and 3 <= len(check) <= 80
        and "|" not in check
        and "•" not in check
        and "@" not in check
        and not check.startswith("http")
    )


def _line_is_role_or_degree_row(stripped: str, current_section: str) -> bool:
    """
    Experience/education entry line: has | (dates) and/or Role — Company / Degree — School.
    Models often emit ### headings instead of the pipe form; those must still become job_title rows.
    """
    if current_section not in _EXPERIENCE_SECTIONS:
        return False
    if "•" in stripped:
        return False
    clean = _clean_line(stripped)
    if not clean or len(clean) < 6:
        return False
    if "|" in stripped:
        return True
    if "\u2014" in clean or " -- " in clean:
        return True
    # En dash or hyphen as separator: "Role - Company"
    if re.search(r"\s[-\u2013]\s", clean):
        return True
    return False


def _parse_resume(text: str) -> list[tuple[str, object]]:
    lines = text.splitlines()
    result: list[tuple[str, object]] = []
    header_block = True
    name_found = False
    title_found = False
    current_section = ""

    for raw in lines:
        stripped = raw.strip()

        if not stripped:
            result.append(("empty", ""))
            continue

        if re.match(r"^[-*=]{3,}$", stripped):
            if header_block and stripped == "---":
                result.append(("hr", ""))
            continue

        if name_found and _is_section_header(stripped):
            clean = _clean_line(stripped)
            header_block = False
            current_section = clean.upper()
            result.append(("section_header", clean))
            continue

        if header_block:
            if not name_found:
                result.append(("name", _clean_line(stripped)))
                name_found = True
                continue
            # Treat as contact when the line clearly contains contact fields
            # (pipes, email, phone, or a URL). Otherwise treat it as an optional
            # title line, which is kept for backward compatibility.
            looks_like_contact = (
                "|" in stripped
                or "@" in stripped
                or re.search(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", stripped)
                or "linkedin" in stripped.lower()
                or stripped.lower().startswith(("http://", "https://"))
            )
            if looks_like_contact:
                result.append(("contact", stripped))
            elif not title_found:
                result.append(("title", _clean_line(stripped)))
                title_found = True
            else:
                result.append(("contact", stripped))
            continue

        if _line_is_role_or_degree_row(stripped, current_section):
            result.append(("job_title", _clean_line(stripped)))
            continue

        skills_section = any(kw in current_section for kw in ("SKILL", "COMPETENC"))
        if skills_section:
            skill_line = stripped
            bulleted = False
            if stripped[0] in ("•", "·", "-", "*") and len(stripped) > 1 and stripped[1] in (" ", "\t"):
                skill_line = stripped[2:].strip()
                bulleted = True
            if ":" in skill_line and not skill_line.startswith("http"):
                idx = skill_line.index(":")
                label = skill_line[:idx].replace("**", "").replace("__", "").strip()
                values = skill_line[idx + 1 :].strip()
                result.append(("skill", (label, values, bulleted)))
                continue

        if stripped[0] in ("•", "·", "-") and (len(stripped) < 2 or stripped[1] in (" ", "\t")):
            result.append(("bullet", stripped.lstrip("•·- ").strip()))
            continue
        if stripped[0] == "*" and len(stripped) > 1 and stripped[1] == " ":
            result.append(("bullet", stripped[2:].strip()))
            continue

        result.append(("body", _clean_line(stripped)))

    return result


def _section_bucket(header_line: str) -> int:
    """
    Canonical section order for DOCX output (lower sorts earlier).
    10 Summary, 20 Skills, 30 Experience/Projects, 40 Education, 45 certs/etc., 50 other.
    """
    u = str(header_line or "").strip().upper()
    if u in (
        "SUMMARY",
        "MARKET TITLE",
        "PROFESSIONAL SUMMARY",
        "OBJECTIVE",
    ):
        return 10
    if "SKILL" in u or "COMPETENC" in u:
        return 20
    if u in (
        "WORK EXPERIENCE",
        "PROFESSIONAL EXPERIENCE",
        "EXPERIENCE",
        "PROJECTS",
        "NOTABLE PROJECTS",
        "KEY PROJECTS",
    ):
        return 30
    if u == "EDUCATION":
        return 40
    if u in (
        "CERTIFICATIONS",
        "CERTIFICATIONS OR ACHIEVEMENTS",
        "ACHIEVEMENTS",
        "AWARDS",
        "ADDITIONAL INFORMATION",
    ):
        return 45
    if u in ("APPLICATION QUESTIONS", "APPLICATION Q&A", "APPLICATION ANSWERS"):
        return 55
    return 50


def _reorder_section_blocks(items: list[tuple[str, object]]) -> list[tuple[str, object]]:
    """Place Summary, Skills, Experience, Education (then other sections) regardless of model order."""
    if not items:
        return items
    i = 0
    preamble: list[tuple[str, object]] = []
    while i < len(items) and items[i][0] != "section_header":
        preamble.append(items[i])
        i += 1

    buckets: dict[int, list[tuple[str, object]]] = {
        10: [],
        20: [],
        30: [],
        40: [],
        45: [],
        50: [],
        55: [],
    }
    while i < len(items):
        if items[i][0] != "section_header":
            preamble.append(items[i])
            i += 1
            continue
        block: list[tuple[str, object]] = [items[i]]
        hdr = str(items[i][1]).strip()
        i += 1
        while i < len(items) and items[i][0] != "section_header":
            block.append(items[i])
            i += 1
        buckets[_section_bucket(hdr)].extend(block)

    out: list[tuple[str, object]] = list(preamble)
    for k in (10, 20, 30, 40, 45, 50, 55):
        out.extend(buckets.get(k, []))
    return out


_SKILL_LABEL_SHORT: dict[str, str] = {
    "cloud, infrastructure & tools": "Cloud & Tools",
    "cloud infrastructure & tools": "Cloud & Tools",
    "cloud & infrastructure": "Cloud & Tools",
    "cloud and infrastructure": "Cloud & Tools",
    "concepts & methodologies": "Engineering Practices",
    "concepts and methodologies": "Engineering Practices",
    "testing, quality & sdlc": "Testing & SDLC",
    "testing quality & sdlc": "Testing & SDLC",
    "databases & data": "Databases",
    "apis & integration": "APIs & Integration",
}


def _short_skill_category_label(label: str) -> str:
    key = str(label or "").strip().lower()
    return _SKILL_LABEL_SHORT.get(key, str(label or "").strip())


def _add_md_runs(paragraph, text: str, base_size_pt: float, bold_base: bool = False) -> None:
    from docx.shared import Pt

    base_size_pt = float(base_size_pt) if base_size_pt else _BODY_PT
    sanitized = _sanitize_text(text)
    parts = re.split(r"(\*\*[^*]+\*\*)", sanitized)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part.replace("**", ""))
            run.bold = bold_base
        run.font.name = _BODY_FONT
        run.font.size = Pt(base_size_pt)


def _build_docx(items: list[tuple[str, object]]) -> object:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Experience / education meta lines: strong contrast for print + screen (not too light).
    _EXP_COLOR_PRIMARY = RGBColor(0x21, 0x21, 0x21)
    _EXP_COLOR_SECONDARY = RGBColor(0x42, 0x42, 0x42)

    for section in doc.sections:
        m = Inches(_MARGINS_IN)
        section.top_margin = m
        section.bottom_margin = m
        section.left_margin = m
        section.right_margin = m

    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(_BODY_PT)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = _LINE_SPACING

    def _para(
        text="",
        bold=False,
        italic=False,
        size=_BODY_PT,
        space_before=0,
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = _LINE_SPACING
        if align is not None:
            p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.name = _BODY_FONT
            run.font.size = Pt(size)
        return p

    def _add_bottom_rule(paragraph) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D4DDE5")

    def _next_non_empty_type(start_idx: int) -> str:
        for next_idx in range(start_idx + 1, len(items)):
            next_type = items[next_idx][0]
            if next_type != "empty":
                return next_type
        return ""

    last_bullet_indices: set[int] = set()
    for _i, (_it, _) in enumerate(items):
        if _it == "bullet":
            for _j in range(_i + 1, len(items)):
                if items[_j][0] != "empty":
                    if items[_j][0] != "bullet":
                        last_bullet_indices.add(_i)
                    break
            else:
                last_bullet_indices.add(_i)

    def _set_right_tab(paragraph, pos: int = 10080) -> None:
        pPr = paragraph._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:pos"), str(pos))
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

    def _add_plain(
        paragraph,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: float = _BODY_PT,
        color=None,
    ) -> None:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = _BODY_FONT
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color

    current_section = ""

    for idx, (item_type, content) in enumerate(items):
        if item_type == "empty":
            continue

        if item_type == "name":
            p = _para(space_before=0, space_after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(str(content))
            run.bold = True
            run.font.name = _BODY_FONT
            run.font.size = Pt(_NAME_PT)

        elif item_type == "title":
            p = _para(space_before=0, space_after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(str(content))
            run.bold = False
            run.font.name = _BODY_FONT
            run.font.size = Pt(_SECTION_TITLE_PT)

        elif item_type == "contact":
            p = _para(space_before=0, space_after=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_md_runs(p, str(content), base_size_pt=_BODY_PT)
            if _next_non_empty_type(idx) == "section_header":
                _add_bottom_rule(p)

        elif item_type == "hr":
            continue

        elif item_type == "section_header":
            display = str(content).strip().title()
            current_section = display.upper()
            p = _para(space_before=14, space_after=8)
            run = p.add_run(display)
            run.bold = True
            run.font.name = _BODY_FONT
            run.font.size = Pt(_SECTION_TITLE_PT)
            _add_bottom_rule(p)

        elif item_type == "job_title":
            raw = re.sub(r"\*\*", "", str(content)).strip()
            # Some upstream resumes mix ``|`` with tabs as the separator
            # (e.g., ``Role  |  Company\tDates  |  Location``). Normalize
            # any run of tab or multi-space that sits between `|` groups
            # so the row splits cleanly into its 4 fields.
            normalized = re.sub(r"\t+", " | ", raw)
            parts = [x.strip() for x in normalized.split("|") if x.strip()]
            is_education = "EDUCATION" in current_section

            # Pipe row from model: Role | Company | Dates | Location (Experience) or
            # School | Degree | Years | Location (Education).
            #
            # Experience (ATS-friendly scan):
            #   Line 1: **Title** | *Company*
            #   Line 2: Location | Dates  (lighter)
            #
            # Education (parallel hierarchy):
            #   Line 1: **Degree** | *School*
            #   Line 2: Location | Years  (lighter)
            if is_education and len(parts) >= 3:
                school = parts[0]
                degree = parts[1] if len(parts) > 1 else ""
                years = parts[2] if len(parts) > 2 else ""
                location = parts[3] if len(parts) > 3 else ""

                p1 = _para(space_before=8, space_after=4)
                if degree and school:
                    _add_plain(p1, degree, bold=True, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)
                    _add_plain(p1, " | ", bold=False, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)
                    _add_plain(
                        p1,
                        school,
                        bold=True,
                        italic=False,
                        size=_BODY_PT,
                        color=_EXP_COLOR_PRIMARY,
                    )
                elif school:
                    _add_plain(p1, school, bold=True, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)
                elif degree:
                    _add_plain(p1, degree, bold=True, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)

                line2 = " | ".join(x for x in (location, years) if x)
                if line2:
                    p2 = _para(space_before=1, space_after=10)
                    _add_plain(
                        p2,
                        line2,
                        bold=False,
                        size=_BODY_PT,
                        color=_EXP_COLOR_SECONDARY,
                    )

            elif not is_education and len(parts) >= 3:
                role = parts[0]
                company = parts[1]
                dates = parts[2] if len(parts) > 2 else ""
                location = parts[3] if len(parts) > 3 else ""

                # Line 1–2: readable gap between title|company and location|dates;
                # line 2: extra space after before bullets.
                p1 = _para(space_before=8, space_after=4)
                if role and company:
                    _add_plain(p1, role, bold=True, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)
                    _add_plain(p1, " | ", bold=False, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)
                    _add_plain(
                        p1,
                        company,
                        bold=True,
                        italic=False,
                        size=_BODY_PT,
                        color=_EXP_COLOR_PRIMARY,
                    )
                elif role:
                    _add_plain(p1, role, bold=True, size=_BODY_PT, color=_EXP_COLOR_PRIMARY)
                elif company:
                    _add_plain(
                        p1,
                        company,
                        bold=True,
                        italic=False,
                        size=_BODY_PT,
                        color=_EXP_COLOR_PRIMARY,
                    )

                line2 = " | ".join(x for x in (location, dates) if x)
                if line2:
                    p2 = _para(space_before=1, space_after=10)
                    _add_plain(
                        p2,
                        line2,
                        bold=False,
                        size=_BODY_PT,
                        color=_EXP_COLOR_SECONDARY,
                    )

            else:
                # Fallback for oddly shaped lines (em-dash, 2 parts, etc.).
                first = parts[0] if parts else raw
                if "\u2014" in first:
                    sub = [x.strip() for x in first.split("\u2014", 1)]
                    role_co = f"{sub[0]} \u2014 {sub[1]}"
                    right = ""
                elif "--" in first:
                    sub = [x.strip() for x in first.split("--", 1)]
                    role_co = f"{sub[0]} \u2014 {sub[1]}"
                    right = ""
                elif len(parts) >= 2:
                    role_co = f"{parts[1]}, {parts[0]}"
                    right = ""
                else:
                    role_co = first
                    right = ""
                p = _para(space_before=6, space_after=4)
                _set_right_tab(p)
                _add_plain(p, role_co, bold=True, size=_BODY_PT)
                if right:
                    _add_plain(p, "\t", size=_BODY_PT)
                    _add_plain(p, right, size=_BODY_PT)

        elif item_type == "bullet":
            is_last_bullet = idx in last_bullet_indices
            in_exp = current_section in _EXPERIENCE_ROLE_SECTIONS
            in_edu = current_section == "EDUCATION"
            if in_exp:
                space_before = 2
                space_after = 14 if is_last_bullet else 6
            elif in_edu:
                space_before = 2
                space_after = 11 if is_last_bullet else 5
            else:
                space_before = 2
                space_after = 9 if is_last_bullet else 5
            p = _para(space_before=space_before, space_after=space_after)
            p.paragraph_format.left_indent = Inches(_BULLET_LEFT_INDENT_IN)
            p.paragraph_format.first_line_indent = Inches(_BULLET_FIRST_LINE_INDENT_IN)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = _LINE_SPACING
            run = p.add_run(_BULLET_LEADER)
            run.font.name = _BODY_FONT
            run.font.size = Pt(_BODY_PT)
            _add_md_runs(p, str(content), base_size_pt=_BODY_PT)

        elif item_type == "skill":
            label, values, _bulleted = content  # type: ignore[misc]
            label_out = _short_skill_category_label(str(label))
            plain_values = _strip_md_spans_for_skill_values(str(values))
            p = _para(space_before=1, space_after=4)
            # Hanging-style wrap for long value lists: continuation lines align under the value
            # column (after tab), not under the category label.
            p.paragraph_format.tab_stops.add_tab_stop(
                Inches(_SKILL_VALUE_TAB_IN),
                WD_TAB_ALIGNMENT.LEFT,
            )
            r_label = p.add_run(f"{label_out}:\t")
            r_label.bold = True
            r_label.font.name = _BODY_FONT
            r_label.font.size = Pt(_BODY_PT)
            # Values stay plain (no per-item bolding); only the category label is bold.
            _add_md_runs(p, plain_values, base_size_pt=_BODY_PT)

        elif item_type == "body":
            sa = 6 if current_section == "SUMMARY" else 0
            p = _para(space_before=0, space_after=sa)
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            _add_md_runs(p, str(content), base_size_pt=_BODY_PT)

    return doc


def build_formatted_resume_docx(resume_text: str, job, profile_name: str) -> tuple[BytesIO, str]:
    """Build resume .docx in memory. Returns (buffer, filename)."""
    items = _reorder_section_blocks(_parse_resume(str(resume_text or "")))
    doc = _build_docx(items)

    candidate_name = (profile_name or "").strip() or "Candidate"
    company_name = getattr(job, "company_name", None) or ""
    job_title = getattr(job, "title", None) or "Role"
    today = date.today().isoformat()
    filename = (
        f"{_safe_filename(candidate_name)}_{_safe_filename(company_name)}"
        f"_{_safe_filename(job_title)}_{today}.docx"
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, filename
